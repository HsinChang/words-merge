from docx import Document


def combine_word_files(file1, file2):
    doc1 = Document(file1)
    doc2 = Document(file2)

    # Select specific chapters from each document
    selected_chapters_doc1 = [doc1.paragraphs[0], doc1.paragraphs[100]]
    selected_chapters_doc2 = [doc2.paragraphs[101], doc2.paragraphs[200]]

    # Create a new document and add the selected chapters from each document
    combined_doc = Document()
    for chapter in selected_chapters_doc1:
        combined_doc.add_paragraph(chapter.text)
    for chapter in selected_chapters_doc2:
        combined_doc.add_paragraph(chapter.text)

    # Save the combined document
    combined_doc.save('combined.docx')

combine_word_files('docx-dump/受理局指南.docx','docx-dump/受理局指南.docx')